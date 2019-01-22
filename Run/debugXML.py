import sys
from docx import Document
import logging
import PortSerial
import time
import xml.etree.ElementTree as ET


logging.basicConfig(format=u'%(levelname)-8s [%(asctime)s] %(message)s', level=logging.DEBUG, filename=u'mylog.log')




def report():

    path_report = "C:\\work\\Test\\MyWPF\\Run"  # str(sys.argv[1])
    port = "COM48"  # str(sys.argv[2])
    setup_id = "1"   # str(sys.argv[3])
    test_name = "Enable"  # str(sys.argv[4])

    tree = ET.parse(path_report + "\\test" + setup_id + ".xml")
    print("Path: " + path_report + "\\test" + setup_id + ".xml")
    root = tree.getroot()

    count = 1
    items = ((0, '', '', ""),)

    if root.get('SetupId') == setup_id:  # Identification that test match to the setup

        for root_group in root:

            if root_group.get("TestId") == test_name:  # Match the test into the setup

                for test_root in root_group:

                    step = test_root.get('Name')
                    expect = test_root.get('Expect')
                    result = test_root.text

                    # print(step)
                    # print(expect)
                    # print(result)

                    items = items + ((count, step, expect, result), )
                    count += 1

    document = Document()
    document.add_heading('Test: Enable Setup', 0)
    p = document.add_paragraph('System information: \n')
    PortSerial.port("\\1", port, 1)
    time.sleep(1)
    p.add_run("Axis info #1:").bold = True
    p.add_run(PortSerial.port("info", port, 0)).bold = True

    table = document.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step Id'
    hdr_cells[1].text = 'Step Description'
    hdr_cells[2].text = 'Expected Result'
    hdr_cells[3].text = 'Results'

    for id, step, des, res in items:
        print(str(id))
        print(str(step))
        print(des)
        print(res)

        row_cells = table.add_row().cells
        row_cells[0].text = str(id)
        row_cells[1].text = str(step)
        row_cells[2].text = des
        row_cells[3].text = res

    document.add_page_break()
    document.save(path_report + "\\Report.docx")

report()


